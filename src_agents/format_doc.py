import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict
from datetime import datetime

class DocumentGenerator:
    """Generate professional Word documents"""
    
    def generate_docx(self, content: str, topic: str, output_path: str = None) -> str:
        """
        Generate Word document from markdown content.
        
        Args:
            content: Markdown content
            topic: Research topic (for title)
            output_path: Where to save the file
            
        Returns:
            Path to generated document
        """
        
        if output_path is None:
            if os.path.exists("research_docs"):
                output_path = f"research_docs/{topic.replace(' ', '_')}_{datetime.now().strftime('%d_%m_%Y_%H_%M_%S')}.docx"
            else:
                output_path = f"{topic.replace(' ', '_')}_{datetime.now().strftime('%d_%m_%Y_%H_%M_%S')}.docx"
        
        # Create document
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Add title
        title = doc.add_heading(topic, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        date_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacing
        
        # Parse markdown and add to document
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            
            if not line:
                continue
            
            # Heading level 1 (# )
            if line.startswith('# '):
                doc.add_heading(line[2:], 1)
            
            # Heading level 2 (## )
            elif line.startswith('## '):
                doc.add_heading(line[3:], 2)
            
            # Heading level 3 (### )
            elif line.startswith('### '):
                doc.add_heading(line[4:], 3)
            
            # Regular paragraph
            else:
                doc.add_paragraph(line)
        
        # Save
        doc.save(output_path)
        
        print(f"  Saved Word document: {output_path}")
        
        return output_path